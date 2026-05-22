"""Конвертация H6_full_report.md → H6_full_report.docx через python-docx.

Поддерживает: заголовки 1–4 уровней (`#` … `####`), параграфы, маркированные
списки (`- `), нумерованные списки (`1. `, `2. `), таблицы, блочный код
(```), inline-italics / bold / code.

Формулы LaTeX оставляются как plain text (Word их не отрендерит, но они
останутся читабельными).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "output" / "h6_verification" / "H6_full_report.md"
DST = ROOT / "output" / "h6_verification" / "H6_full_report.docx"


def parse_md(text: str) -> list[dict]:
    """Простой парсер: возвращает список блоков-словарей."""
    blocks = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        if not ln.strip():
            i += 1
            continue
        # Heading
        m = re.match(r"^(#{1,5})\s+(.*)", ln)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)),
                            "text": m.group(2).strip()})
            i += 1
            continue
        # Code block
        if ln.startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].startswith("```"):
                buf.append(lines[j])
                j += 1
            blocks.append({"type": "code", "text": "\n".join(buf)})
            i = j + 1
            continue
        # Table (markdown table)
        if ln.lstrip().startswith("|") and i + 1 < len(lines) and \
           re.match(r"^\s*\|[\s\-:\|]+\|\s*$", lines[i + 1]):
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            # Parse rows
            rows = []
            for r in tbl:
                if re.match(r"^\s*\|[\s\-:\|]+\|\s*$", r):
                    continue
                cells = [c.strip() for c in r.strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue
        # Bulleted list
        if re.match(r"^[-*]\s+", ln):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(re.sub(r"^[-*]\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "bullet", "items": items})
            continue
        # Numbered list
        if re.match(r"^\d+\.\s+", ln):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "ordered", "items": items})
            continue
        # Plain paragraph (collect until blank)
        buf = [ln]
        i += 1
        while i < len(lines) and lines[i].strip() and \
              not re.match(r"^(#{1,5})\s+", lines[i]) and \
              not lines[i].lstrip().startswith("|") and \
              not lines[i].startswith("```") and \
              not re.match(r"^[-*]\s+", lines[i]) and \
              not re.match(r"^\d+\.\s+", lines[i]):
            buf.append(lines[i])
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(buf)})
    return blocks


def add_inline(par, text: str):
    """Поддержка **bold**, *italic*, `code`, $math$ → как plain run."""
    # Очень простой парсер: проходим токенами **…**, *…*, `…`, $…$.
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group()
        if tok.startswith("**"):
            run = par.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith("*"):
            run = par.add_run(tok[1:-1]); run.italic = True
        elif tok.startswith("`"):
            run = par.add_run(tok[1:-1])
            run.font.name = "Courier New"
        elif tok.startswith("$"):
            run = par.add_run(tok[1:-1]); run.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def build_docx(blocks: list[dict], dst: Path):
    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for b in blocks:
        if b["type"] == "heading":
            h = doc.add_heading(b["text"], level=min(b["level"], 4))
        elif b["type"] == "paragraph":
            p = doc.add_paragraph()
            add_inline(p, b["text"])
        elif b["type"] == "bullet":
            for it in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, it)
        elif b["type"] == "ordered":
            for it in b["items"]:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, it)
        elif b["type"] == "code":
            p = doc.add_paragraph()
            run = p.add_run(b["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif b["type"] == "table":
            rows = b["rows"]
            if not rows:
                continue
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for r_idx, r in enumerate(rows):
                for c_idx, cell in enumerate(r):
                    if c_idx < len(t.rows[r_idx].cells):
                        para = t.rows[r_idx].cells[c_idx].paragraphs[0]
                        para.clear()
                        add_inline(para, cell)
                        # Make header row bold
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True

    doc.save(dst)
    print(f"Wrote {dst}")


def main():
    text = SRC.read_text(encoding="utf-8")
    blocks = parse_md(text)
    build_docx(blocks, DST)


if __name__ == "__main__":
    main()
