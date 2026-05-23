"""
Build HSE SPB SOEM term-paper .docx template.

Specs per «Guidelines for preparing the term paper» (SPB SOEM, 2023):
- Times New Roman 12pt body, 1.5 line spacing (NOT 14pt; guideline §6.1)
- A4 page; margins: left 35mm, right ≥10mm, top/bottom ≥20mm
- Page numbers bottom (center or right), starting from page 2
- Main text justified; headings centered
- Paragraph indent 1.25cm
- Heading 1/2/3 styles at TNR 12pt bold (matching main text)
- Reference list 12pt / 1.5 spacing
- Tables, figures, footnotes 10pt / 1.0 spacing

Usage:
    python3 coursework/build_template.py

Produces coursework/hse_template.docx — pass to pandoc via
    pandoc ... --reference-doc=coursework/hse_template.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_default_font(doc, name='Times New Roman', size_pt=12):
    style = doc.styles['Normal']
    style.font.name = name
    # Set East-Asian and complex-script font names too (Word stores them separately)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), name)
    style.font.size = Pt(size_pt)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(1.25)
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WAP
    style.paragraph_format.alignment = WAP.JUSTIFY


def set_page_size_and_margins(doc):
    """Per HSE SPB SOEM guideline §6.1: left 35mm, right ≥10mm, top/bottom ≥20mm."""
    for section in doc.sections:
        section.page_height = Mm(297)  # A4
        section.page_width = Mm(210)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(35)
        section.right_margin = Mm(10)


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # PAGE field
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)
        instr = OxmlElement('w:instrText')
        instr.text = 'PAGE'
        run._r.append(instr)
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)


def style_headings(doc):
    """Headings — TNR 12pt bold, centered, per guideline §6.1."""
    for level in (1, 2, 3):
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.italic = False
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), 'Times New Roman')
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def style_footnote(doc):
    """Footnotes — TNR 10pt single-spaced per §6.1."""
    if 'Footnote Text' in doc.styles:
        style = doc.styles['Footnote Text']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), 'Times New Roman')
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def main():
    doc = Document()
    set_default_font(doc)
    set_page_size_and_margins(doc)
    add_page_number_footer(doc)
    style_headings(doc)
    style_footnote(doc)

    # Sample content so styles are visible if anyone opens the file directly
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Cm(0)
    r = p_title.add_run(
        'HSE SPB SOEM Term Paper Template — TNR 12pt, line spacing 1.5, A4, '
        'left margin 35mm, right 10mm, top/bottom 20mm'
    )
    r.bold = True

    doc.add_paragraph()
    doc.add_heading('Глава 1. Образец заголовка первого уровня', level=1)
    doc.add_paragraph(
        'Это пример основного текста: Times New Roman, 12 пт, межстрочный '
        'интервал 1,5, отступ первой строки 1,25 см, выравнивание по ширине. '
        'Шаблон используется как --reference-doc при сборке итогового документа '
        'через pandoc.'
    )

    doc.add_heading('1.1. Заголовок второго уровня', level=2)
    doc.add_paragraph(
        'Раздел второго уровня для разделения главы на подсекции. Нумерация '
        '— вручную (1.1, 1.2, ...) либо через auto-numbering в Word.'
    )

    doc.add_heading('1.1.1. Заголовок третьего уровня', level=3)
    doc.add_paragraph('Самый глубокий уровень, типичный для разделов методов.')

    out = 'coursework/hse_template.docx'
    doc.save(out)
    print(f'Saved {out}')


if __name__ == '__main__':
    main()
